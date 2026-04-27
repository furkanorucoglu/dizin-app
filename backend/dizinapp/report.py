"""PDF/DOCX export helpers for Dizin App.

This module intentionally avoids WeasyPrint so local macOS installs do not need
GTK/Pango/GObject system libraries. PDF output is generated with ReportLab;
Word output is generated with python-docx.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable


def _value(obj: Any, key: str, default: Any = None) -> Any:
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def _text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _entry_no(entry: Any) -> str:
    value = _value(entry, "paragraph_index", _value(entry, "id", ""))
    return _text(value)


def _headword(entry: Any) -> str:
    return _text(_value(entry, "headword", "")) or "-"


def _raw_original_line(entry: Any) -> str:
    raw = _text(_value(entry, "raw_text", ""))
    if raw:
        return raw
    return f"{_headword(entry)} {_format_pages(_value(entry, 'original_pages', []))}".strip()


def _confidence_label(entry: Any) -> str:
    value = _text(_value(entry, "confidence", "")).lower()
    return {"high": "Yüksek", "medium": "Orta", "low": "Düşük"}.get(value, value or "-")


def _manual_label(entry: Any) -> str:
    return "Manuel düzeltildi" if bool(_value(entry, "manually_edited", False)) else "Otomatik"


def _page_part(page_ref: Any) -> str:
    raw = _text(_value(page_ref, "raw", ""))
    if raw:
        return raw
    start = _value(page_ref, "start", None)
    end = _value(page_ref, "end", None)
    if start is None and end is None:
        return ""
    if end is None or end == start:
        return str(start)
    return f"{start}–{end}"


def _format_pages(page_refs: Iterable[Any] | None) -> str:
    if not page_refs:
        return "-"
    parts = [_page_part(p) for p in page_refs]
    parts = [p for p in parts if p]
    return ", ".join(parts) if parts else "-"


def _original_pages(entry: Any) -> str:
    return _format_pages(_value(entry, "original_pages", []))


def _translated_pages(entry: Any) -> str:
    return _format_pages(_value(entry, "translated_pages", []))


def _safe_title(title: str | None) -> str:
    return _text(title) or "Dizin"


def _register_fonts():
    """Register a Unicode font when available and return font names."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates_regular = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/local/share/fonts/dejavu/DejaVuSans.ttf",
    ]
    candidates_bold = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/local/share/fonts/dejavu/DejaVuSans-Bold.ttf",
    ]

    regular = next((p for p in candidates_regular if Path(p).exists()), None)
    bold = next((p for p in candidates_bold if Path(p).exists()), None)

    if regular:
        pdfmetrics.registerFont(TTFont("DizinSans", regular))
        if bold:
            pdfmetrics.registerFont(TTFont("DizinSans-Bold", bold))
        else:
            pdfmetrics.registerFont(TTFont("DizinSans-Bold", regular))
        return "DizinSans", "DizinSans-Bold"

    # Fallback can render ASCII/Latin-1; Unicode font is preferred on macOS/Linux.
    return "Helvetica", "Helvetica-Bold"


def _pdf_styles(base_font: str, bold_font: str):
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="DizinTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="DizinSubtitle",
        parent=styles["BodyText"],
        fontName=base_font,
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=14,
    ))
    styles.add(ParagraphStyle(
        name="DizinCell",
        parent=styles["BodyText"],
        fontName=base_font,
        fontSize=8.5,
        leading=11,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#1e293b"),
    ))
    styles.add(ParagraphStyle(
        name="DizinCellBold",
        parent=styles["BodyText"],
        fontName=bold_font,
        fontSize=8.5,
        leading=11,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#0f172a"),
    ))
    styles.add(ParagraphStyle(
        name="DizinSmall",
        parent=styles["BodyText"],
        fontName=base_font,
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor("#64748b"),
    ))
    return styles


def _p(styles: Any, text: Any, style_name: str = "DizinCell"):
    from xml.sax.saxutils import escape
    from reportlab.platypus import Paragraph

    return Paragraph(escape(_text(text)), styles[style_name])


def _summary(entries: list[Any]) -> dict[str, int]:
    return {
        "total": len(entries),
        "manual": sum(1 for e in entries if bool(_value(e, "manually_edited", False))),
        "low": sum(1 for e in entries if _text(_value(e, "confidence", "")).lower() == "low"),
        "medium": sum(1 for e in entries if _text(_value(e, "confidence", "")).lower() == "medium"),
        "high": sum(1 for e in entries if _text(_value(e, "confidence", "")).lower() == "high"),
    }


def generate_index_pdf(title: str | None, entries: Iterable[Any]) -> bytes:
    """Generate the final Turkish index as a readable PDF."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle

    rows = list(entries)
    buffer = BytesIO()
    base_font, bold_font = _register_fonts()
    styles = _pdf_styles(base_font, bold_font)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.25 * cm,
        leftMargin=1.25 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
        title=f"{_safe_title(title)} - Dizin",
    )

    s = _summary(rows)
    story = [
        _p(styles, f"{_safe_title(title)} - Türkçe Dizin", "DizinTitle"),
        _p(
            styles,
            f"Toplam {s['total']} giriş · Manuel düzeltme {s['manual']} · Yüksek/Orta/Düşük güven: {s['high']}/{s['medium']}/{s['low']} · {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            "DizinSubtitle",
        ),
        Spacer(1, 8),
    ]

    data = [[
        _p(styles, "#", "DizinCellBold"),
        _p(styles, "Dizin kelimesi", "DizinCellBold"),
        _p(styles, "Türkçe sayfalar", "DizinCellBold"),
        _p(styles, "Güven", "DizinCellBold"),
        _p(styles, "Durum", "DizinCellBold"),
    ]]

    for entry in rows:
        data.append([
            _p(styles, _entry_no(entry), "DizinSmall"),
            _p(styles, _headword(entry), "DizinCellBold"),
            _p(styles, _translated_pages(entry)),
            _p(styles, _confidence_label(entry)),
            _p(styles, _manual_label(entry), "DizinSmall"),
        ])

    table = Table(data, colWidths=[1.0 * cm, 6.4 * cm, 5.0 * cm, 2.0 * cm, 3.0 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(table)
    doc.build(story)
    return buffer.getvalue()


def generate_comparison_pdf(title: str | None, entries: Iterable[Any], *args: Any, **kwargs: Any) -> bytes:
    """Generate a side-by-side PDF comparing the first index with final index.

    Extra *args/**kwargs are accepted for backward compatibility with older route code
    that passed PDF paths/offsets to this function.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle

    rows = list(entries)
    buffer = BytesIO()
    base_font, bold_font = _register_fonts()
    styles = _pdf_styles(base_font, bold_font)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=1.0 * cm,
        leftMargin=1.0 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
        title=f"{_safe_title(title)} - Karşılaştırma",
    )

    s = _summary(rows)
    story = [
        _p(styles, f"{_safe_title(title)} - Dizin Karşılaştırma", "DizinTitle"),
        _p(
            styles,
            f"Solda ilk verilen dizin satırı, sağda son düzeltilmiş Türkçe dizin yer alır. Toplam {s['total']} giriş · Manuel düzeltme {s['manual']} · {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            "DizinSubtitle",
        ),
        Spacer(1, 8),
    ]

    data = [[
        _p(styles, "#", "DizinCellBold"),
        _p(styles, "İlk verilen dizin dosyası", "DizinCellBold"),
        _p(styles, "Son düzeltilmiş Türkçe dizin", "DizinCellBold"),
        _p(styles, "Kontrol", "DizinCellBold"),
    ]]

    for entry in rows:
        data.append([
            _p(styles, _entry_no(entry), "DizinSmall"),
            _p(styles, _raw_original_line(entry)),
            _p(styles, f"{_headword(entry)} {_translated_pages(entry)}"),
            _p(styles, f"{_confidence_label(entry)} · {_manual_label(entry)}", "DizinSmall"),
        ])

    table = Table(data, colWidths=[1.0 * cm, 11.2 * cm, 11.2 * cm, 3.0 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2ff")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ]))
    story.append(table)
    doc.build(story)
    return buffer.getvalue()


def generate_index_docx(title: str | None, entries: Iterable[Any]) -> bytes:
    """Generate the final Turkish index as a clean Word document.

    The Word export intentionally avoids the dashboard/report table layout.
    It follows the user's sample DIZIN_TR.docx style: a simple title and
    one index paragraph per entry, so the file can be edited like a normal
    book index in Word.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    rows = list(entries)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(12)
    title_run = title_paragraph.add_run("DİZİN")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    # Keep the export visually close to the provided Word sample: no summary,
    # no confidence/status columns, no table borders. Each line is editable text.
    for entry in rows:
        headword = _headword(entry).strip()
        pages = _translated_pages(entry).strip()
        if not headword:
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        run = paragraph.add_run(f"{headword} {pages}" if pages else headword)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
